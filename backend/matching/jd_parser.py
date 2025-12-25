# # backend/matching/jd_parser.py

# import os
# import json
# from dotenv import load_dotenv
# from docx import Document
# from PIL import Image
# import google.generativeai as genai
# import fitz

# from matching.vector_store import embed_chunks

# load_dotenv()
# genai.configure(api_key=os.getenv("APP_API_KEY"))

# # =========================
# # File type detection (SAFE – extension based)
# # =========================
# def f_type(path):
#     ext = os.path.splitext(path)[1].lower()
#     if ext == ".pdf":
#         return "pdf"
#     if ext == ".docx":
#         return "docx"
#     if ext in [".png", ".jpg", ".jpeg"]:
#         return "image"
#     if ext == ".txt":
#         return "text"
#     return "unknown"

# # =========================
# # PDF extraction
# # =========================
# def extract_from_pdf(path):
#     doc = fitz.open(path)
#     return "".join(page.get_text() for page in doc)

# # =========================
# # DOCX extraction
# # =========================
# def extract_from_docx(path):
#     doc = Document(path)
#     return "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())

# # =========================
# # Image extraction (Gemini Vision)
# # =========================
# def extract_from_image(path):
#     img = Image.open(path)

#     model = genai.GenerativeModel("models/gemini-1.5-flash")
#     response = model.generate_content([
#         img,
#         "Extract all text from this image and return plain text only."
#     ])

#     return response.text.strip()

# # =========================
# # JD chunking via Gemini
# # =========================
# def chunk_jd(text):
#     model = genai.GenerativeModel("gemini-2.5-flash")

#     response = model.generate_content([
#         text,
#         "Return ONLY a JSON list like "
#         "[{\"chunk_text\": \"...\", \"section\": \"job_description\"}]"
#     ])

#     raw = response.text.strip()
#     if raw.startswith("```"):
#         raw = raw.replace("```json", "").replace("```", "").strip()

#     try:
#         return json.loads(raw)
#     except Exception:
#         return []

# # =========================
# # Main JD processing
# # =========================
# def process_jd_and_embed(file_path, session_id):
#     t = f_type(file_path)

#     if t == "pdf":
#         text = extract_from_pdf(file_path)
#     elif t == "docx":
#         text = extract_from_docx(file_path)
#     elif t == "image":
#         text = extract_from_image(file_path)
#     elif t == "text":
#         with open(file_path, "r", encoding="utf-8") as f:
#             text = f.read()
#     else:
#         raise ValueError("Unsupported JD file type")

#     chunks = chunk_jd(text)
#     embed_chunks(chunks, session_id, "jd")






# backend/matching/jd_parser.py

import os
import json
from dotenv import load_dotenv
from docx import Document
from PIL import Image
from google import genai
import fitz  # PyMuPDF

from matching.vector_store import embed_chunks

load_dotenv()

# =========================
# Gemini client (NEW SDK)
# =========================
def get_gemini_client():
    api_key = os.getenv("APP_API_KEY")
    if not api_key:
        raise RuntimeError("APP_API_KEY not set")
    return genai.Client(api_key=api_key)

# =========================
# File type detection
# =========================
def f_type(path):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return "pdf"
    if ext == ".docx":
        return "docx"
    if ext in [".png", ".jpg", ".jpeg"]:
        return "image"
    if ext == ".txt":
        return "text"
    return "unknown"

# =========================
# PDF extraction
# =========================
def extract_from_pdf(path):
    doc = fitz.open(path)
    return "".join(page.get_text() for page in doc)

# =========================
# DOCX extraction
# =========================
def extract_from_docx(path):
    doc = Document(path)
    return "\n".join(
        p.text.strip() for p in doc.paragraphs if p.text.strip()
    )

# =========================
# Image extraction (Gemini Vision)
# =========================
def extract_from_image(path):
    client = get_gemini_client()
    img = Image.open(path)

    response = client.models.generate_content(
        model="gemini-1.5-flash",
        contents=[
            img,
            "Extract all text from this image. Return ONLY plain text."
        ]
    )

    return response.text.strip()

# =========================
# JD intelligent chunking (Gemini)
# =========================
def chunk_jd(text):
    client = get_gemini_client()

    response = client.models.generate_content(
        model="gemini-1.5-flash",
        contents=[
            text,
            (
                "Split this job description into meaningful chunks. "
                "Return ONLY valid JSON in this format:\n"
                "[{\"chunk_text\": \"...\", \"section\": \"job_description\"}]"
            )
        ]
    )

    raw = response.text.strip()

    # Remove markdown fences if Gemini adds them
    if raw.startswith("```"):
        raw = raw.replace("```json", "").replace("```", "").strip()

    try:
        return json.loads(raw)
    except Exception:
        return []

# =========================
# Main JD processing
# =========================
def process_jd_and_embed(file_path, session_id):
    t = f_type(file_path)

    if t == "pdf":
        text = extract_from_pdf(file_path)
    elif t == "docx":
        text = extract_from_docx(file_path)
    elif t == "image":
        text = extract_from_image(file_path)
    elif t == "text":
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
    else:
        raise ValueError("Unsupported JD file type")

    chunks = chunk_jd(text)
    embed_chunks(chunks, session_id, "jd")
